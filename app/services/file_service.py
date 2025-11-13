import os
import shutil
import subprocess
import sys
import tarfile
import tempfile
from abc import ABC, abstractmethod
from io import BytesIO
from math import ceil
from typing import Any, Optional
from zipfile import ZipFile

import docx  # type: ignore
import pandas as pd
from loguru import logger
from pandas.errors import EmptyDataError
from PyPDF2 import PdfFileReader

from app.core.config import settings
from app.schemas import DataChunk, FileStatus, ObjectContents
from app.services.utils.disk_usage import check_archive_size

ARCHIVE_EXTENSIONS = ('.zip', '.tar', '.tar.gz', '.tar.bz2')
CONTAINER_TYPES = ('.csv', '.doc', '.docx', '.xlsx', '.xls', '.pdf')
CSV_DELIMITERS = (',', '\t', ';', '|')
CSV_ENCODINGS = ('ISO-8859-1', 'utf-8', 'windows-1252', 'utf-16', 'utf-16le', 'utf-16be', 'ascii')

class FileService(ABC):
    @abstractmethod
    async def read_data(self, fetch_path: str, service_client: Optional[Any] = None) -> Optional[bytes]:
        """
        Reads data from a source by fetch path.

        Args:
            service_client: Optional. Client session for the connections from decorator.
            fetch_path: The path of a retrieving object.

        Returns:
            The data as bytes if successful, or None if an error occurs.
        """
        ...

    @staticmethod
    def get_tar_read_mode(value: str) -> str:
        """
        Select read mode for tar archive basing on incoming value.

        Args:
            value: path to a file of extension of a file

        Returns:
            mode: string value for tarfile lib
        """
        if value.endswith('tar.gz'):
            mode = 'r:gz'
        elif value.endswith('tar.bz2'):
            mode = 'r:bz2'
        else:
            mode = 'r'
        return mode

    @staticmethod
    def get_tar_archive_extension(file_name: str) -> str:
        """
        Get an archive type basing on file name.

        Args:
            file_name: name of an object on local storage or at a resource

        Returns:
            archive_type: extension of an archive
        """
        if file_name.endswith('tar.gz'):
            archive_type = 'tar.gz'
        elif file_name.endswith('tar.bz2'):
            archive_type = 'tar.bz2'
        else:
            archive_type = 'tar'
        return archive_type

    async def create_file_chunks(self, fetch_path: str, object_name: str, size: int) -> list[DataChunk]:
        """
        Creation data chunks for GitHubBranch.

        Args:
            fetch_path: str - path for retrieving file information
            object_name: str - name of object
            size: size of an object

        Returns:
            data_chunks: list[DataChunk] or []
        """
        data_chunks: list[DataChunk] = []
        if object_name.endswith(settings.UNSUPPORTED_EXTENSIONS):
            return data_chunks

        if fetch_path.endswith(CONTAINER_TYPES):
            if os.path.exists(fetch_path):
                with open(fetch_path, 'rb') as f:
                    size = self.get_content_size(f.read(), object_name)
            else:
                object_data = await self.read_data(fetch_path=fetch_path)
                if not object_data:
                    return data_chunks
                size = self.get_content_size(object_data, object_name)

        if not size:
            return data_chunks

        for i in range(ceil(size / settings.CHUNK_BYTES_CAPACITY)):
            data_chunks.append(
                DataChunk(  # type: ignore
                    object_name=fetch_path.rsplit('::')[-1],
                    fetch_path=fetch_path,
                    offset=str(i * settings.CHUNK_BYTES_CAPACITY),
                    limit=settings.CHUNK_BYTES_CAPACITY,
                    instance_id=settings.SCANNER_ID,
                )
            )
        return data_chunks

    async def collect_file_chunks(self, content: ObjectContents) -> Optional[ObjectContents]:
        """
        Collecting data chunks for every file object.
         Also set various statues if object is an archive, and it can not be unpacked .
        Args:
            content: metadata for file object

        Returns:
            Optional. Content with set status and chunks.
        """
        if not content.size:
            content.status = FileStatus.SCANNED
            return content

        try:
            if content.fetch_path.endswith(ARCHIVE_EXTENSIONS):
                if content.size >= settings.INITIAL_DISK_SPACE:
                    content.status = FileStatus.SKIPPED
                    return content
                object_data = await self.read_data(fetch_path=content.fetch_path)

                if not object_data:
                    return None

                if not self.check_archive_memory_cost(content.full_path, object_data):
                    logger.info(
                        f'{content.full_path} temporary skipped, will be tried to unpack during the next iteration'
                    )
                    return None

                else:
                    data_chunks: list[DataChunk] = []
                    for fetch_path, object_name in self.unpack_archive_locally(content.full_path, object_data):
                        chunks = await self.create_file_chunks(
                            fetch_path, object_name, os.path.getsize(fetch_path)  # type: ignore
                        )
                        data_chunks.extend(chunks)
                    content.data_chunks = data_chunks
            else:
                content.data_chunks = await self.create_file_chunks(
                    fetch_path=content.fetch_path, object_name=content.object_name, size=content.size
                )
            if not content.data_chunks:
                content.status = FileStatus.SCANNED
            return content
        except Exception as e:
            logger.error(f'Error during collecting chunks for {content.full_path}: {e}')
            content.status = FileStatus.FAILED
        return content

    def get_uncompressed_size(self, full_path: str, object_data: bytes) -> int:
        """
        Recursively calculates the uncompressed size of an archive, including any nested archives.

        Args:
            object_data: The binary data of the archive.
            full_path: The full path or file name of the archive.

        Returns:
            The total uncompressed size of all files within the archive, including nested archives.
        """
        uncompressed_size = 0
        if full_path.endswith('.zip') or object_data.startswith(b'PK'):
            with ZipFile(BytesIO(object_data)) as zip_file:
                for info in zip_file.infolist():
                    uncompressed_size += info.file_size
                    if any(info.filename.endswith(ext) for ext in ARCHIVE_EXTENSIONS):
                        nested_data = zip_file.read(info.filename)
                        uncompressed_size += self.get_uncompressed_size(info.filename, nested_data)

        elif full_path.endswith(('.tar', '.tar.gz', '.tar.bz2')):
            mode = self.get_tar_read_mode(full_path)
            with tarfile.open(fileobj=BytesIO(object_data), mode=mode) as tar_ref:
                for member in tar_ref.getmembers():
                    uncompressed_size += member.size
                    if any(member.name.endswith(ext) for ext in ARCHIVE_EXTENSIONS):
                        nested_data = tar_ref.extractfile(member).read()  # type: ignore
                        uncompressed_size += self.get_uncompressed_size(member.name, nested_data)
        return uncompressed_size

    def check_archive_memory_cost(self, full_path: str, object_data: bytes) -> Optional[int]:  # type: ignore
        """
        Checks if the uncompressed size of an archive file is within acceptable memory limits.
        This method calculates the uncompressed size of an archive (including nested archives).

        Args:
            full_path: The full path or filename of the archive.
            object_data: The binary data of the archive.
        """
        try:
            uncompressed_size = self.get_uncompressed_size(full_path, object_data)
            return check_archive_size(uncompressed_size)
        except Exception as e:
            logger.error(f"Error estimating archive size: {e}")

    def process_nested_objects(self, archive_path: str, object_name: str) -> list[tuple[str, str]]:
        """
        Processes nested objects in generator within an archive. Handles directories, normal files,
        and various archive types. Useful for archives that contain other archives.

        Args:
            archive_path: The directory where the object is located.
            object_name: The name of the file or folder to be processed within the archive.

        Yields:
            Tuple containing the file paths and file names of processed object
            from the nested structure.
        """
        file_path = os.path.join(archive_path, object_name)
        try:
            if os.path.isdir(file_path):
                yield from self.process_directory(file_path)
            elif object_name.endswith('.zip'):
                yield from self.process_archive(file_path, 'zip')
            elif object_name.endswith(('tar', 'tar.gz', 'tar.bz2')):
                archive_type = self.get_tar_archive_extension(object_name)
                yield from self.process_archive(file_path, archive_type)
            else:
                yield file_path, object_name
        except Exception as e:
            logger.error(e)

    def process_archive(
        self, file_path: str, archive_type: str, retry: Optional[bool] = True
    ) -> Optional[tuple[str, str]]:
        """
        Generator for processing archive from unpacked archive. This method supports multiple archive formats,
        including zip, tar, tar.gz, and tar.bz2. After extraction, the method processes the contents of
        the extracted directory, handling any nested structures or files.

        The method first creates a directory for extraction, then extracts the contents based on the specified
        archive type. If an error occurs during extraction, and if 'retry' is True, the method attempts a
        retry for zip files. This is based on a check of the file's signature. After extraction, the method
        processes the extracted directory's contents and returns the paths and names of the processed files.

        Args:
            file_path: The path of the archive file to be processed.
            archive_type: The type of the archive ('zip', 'tar', 'tar.gz', 'tar.bz2').
            retry: Indicates if a retry should be attempted on extraction failure. Default is True.

        Yields:
            Optional[tuple[str, str]]: A list of tuples with the file path and file name of each file extracted
            and processed from the archive.
            Returns None if an error occurs and no retry is attempted or if the retry fails.
        """
        try:
            extract_dir = os.path.join(file_path + "_extracted_archive")
            if not os.path.exists(extract_dir):
                os.makedirs(extract_dir, exist_ok=True)
                if archive_type == 'zip':
                    with ZipFile(file_path, 'r') as archive:
                        archive.extractall(path=extract_dir)
                else:  # tar, tar.gz, tar.bz2
                    mode = self.get_tar_read_mode(archive_type)
                    with tarfile.open(file_path, mode) as archive:
                        archive.extractall(path=extract_dir)
            os.remove(file_path)
            yield from self.process_directory(extract_dir)
        except Exception as e:
            logger.error(e)
            # If tar archive first byte starts with `PK` it means that archive must be opened as ZIP archive
            if retry:
                try:
                    with open(file_path, 'rb') as f:
                        data = f.read()
                    if data.startswith(b'PK'):
                        yield from self.process_archive(file_path, 'zip', False)
                except Exception as inner_e:
                    logger.error(inner_e)

    def process_directory(self, directory: str) -> tuple[str, str]:
        """
        Generator recursively processes each file or directory within the specified directory.
        Handles directories, normal files, and various archive types (zip, tar, tar.gz, tar.bz2).

        Args:
            directory: The path of the directory to be processed.

        Yields:
            tuple which contains the file path and file name of processed file within the directory
        """
        try:
            for nested_object_name in os.listdir(directory):
                nested_path = os.path.join(directory, nested_object_name)
                if os.path.isdir(nested_path):
                    yield from self.process_directory(nested_path)
                elif nested_path.endswith('.zip'):
                    yield from self.process_archive(nested_path, 'zip')
                elif nested_path.endswith(('tar', 'tar.gz', 'tar.bz2')):
                    archive_type = self.get_tar_archive_extension(nested_object_name)
                    yield from self.process_archive(nested_path, archive_type)
                else:
                    yield nested_path, nested_object_name
        except Exception as e:
            logger.error(e)

    def unpack_archive_locally(self, full_path: str, object_data: bytes) -> tuple[str, str]:
        """
        Unpacks an archive to a local directory and processes its contents. First reads the
        archive data, then extracts it, and finally processes the extracted content to retrieve
        file paths and names. This method is the starting point for processing archives.

        Args:
            full_path: The full path with the bucket name and object name in '<bucket>/<object>' format.
            object_data: archive data for unpacking.

        Returns:
            A list of tuples, each containing the file path and file name of every file extracted
            and processed from the archive.
        """
        # added postfix for extracted archive to exclude duplications with bucket directories
        extract_path = os.path.join(settings.LOCAL_STORED_ARCHIVES_PATH, f"{full_path}_extracted_archive")
        try:
            if not os.path.exists(extract_path):
                if full_path.endswith('zip') or object_data.startswith(b'PK'):
                    with ZipFile(BytesIO(object_data)) as zip_file:
                        zip_file.extractall(path=extract_path)
                elif full_path.endswith(('tar', 'tar.gz', 'tar.bz2')):
                    mode = self.get_tar_read_mode(full_path)
                    with tarfile.open(fileobj=BytesIO(object_data), mode=mode) as tar_ref:
                        tar_ref.extractall(path=extract_path)
            del object_data
            for object_name in os.listdir(extract_path):
                yield from self.process_nested_objects(extract_path, object_name)
        except OSError as e:
            logger.error(f"Error during extraction {full_path}: {e}")
            if os.path.exists(extract_path):
                shutil.rmtree(extract_path)
        except Exception as e:
            logger.error(e)

    @staticmethod
    def prepare_file(object_data: bytes, file_name: str, limit: int, offset: int) -> Any:
        """
        Copy of presidio prepare file method for S3 without generator.
        Archives were removed because they will be converted directly.
        """
        try:
            if file_name.endswith('pdf'):
                data = ''
                pdf_reader = PdfFileReader(BytesIO(object_data))
                count = pdf_reader.numPages
                for page_num in range(count):
                    page = pdf_reader.getPage(page_num)
                    data += page.extractText()
                return data[offset : offset + limit]  # type: ignore

            elif file_name.endswith(('xlsx', 'xls')):
                if file_name.endswith('xlsx'):
                    sheets = pd.read_excel(BytesIO(object_data), engine='openpyxl', header=None, sheet_name=None)
                else:
                    sheets = pd.read_excel(BytesIO(object_data), engine='xlrd', header=None, sheet_name=None)
                combined_df = pd.concat(sheets.values(), ignore_index=True)
                return combined_df.iloc[offset : offset + limit]  # type: ignore

            elif file_name.endswith('csv'):
                for delimiter in CSV_DELIMITERS:
                    for encoding in CSV_ENCODINGS:
                        try:
                            content = pd.read_csv(
                                BytesIO(object_data),
                                encoding=encoding,
                                delimiter=delimiter,
                                header=None,
                                engine="python",
                            )
                            return content.iloc[offset : offset + limit]
                        except (UnicodeError, pd.errors.ParserError):
                            continue

            elif file_name.endswith('docx'):
                if file_name.endswith('docx'):
                    doc = docx.Document(BytesIO(object_data))
                    full_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
                    return '\n'.join(full_text)[offset : offset + limit]  # type: ignore

            elif file_name.endswith('doc'):
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=True) as doc_temp_file:
                    doc_temp_file.write(object_data)
                    doc_temp_file.flush()  # Ensure data is written to file

                    # Run Antiword to convert the .doc file to text
                    # The output is captured in the subprocess output

                    result = subprocess.run(
                        ['antiword', doc_temp_file.name], check=True, text=True, capture_output=True
                    )
                    return result.stdout[offset : offset + limit]  # type: ignore

            elif file_name.endswith(settings.UNSUPPORTED_EXTENSIONS):
                return ''

            else:
                return object_data.decode(errors='replace')[offset : offset + limit]

        except EmptyDataError:
            logger.warning(f"No columns to parse from file {file_name}")
            return None
        except Exception as e:
            logger.error(f'File {file_name}. Local Exception: {e}')
        return None

    @staticmethod
    def get_content_size(object_data: bytes, file_name: str) -> int:
        """
        Calculates the content size of various file types given their byte data and file name.
        Supports different file formats like PDF, Excel (xlsx, xls), CSV, Word (doc, docx), and images.
        For non-text files and unsupported formats, it defaults to the byte size of the data.

        The method handles specific parsing for each file type to determine content size. For 'doc' files,
        it includes conversion to 'docx' using LibreOffice. Errors in processing are logged, and the method
        returns the calculated size or zero in case of exceptions.

        Args:
            object_data: The byte data of the file.
            file_name: The name of the file, including its extension.

        Returns:
            content_size: The calculated size of the content based on the file type. Returns zero if an error occurs.
        """
        content_size = 0
        try:
            if file_name.endswith('pdf'):
                pdf_reader = PdfFileReader(BytesIO(object_data))
                count = pdf_reader.numPages
                for page_num in range(count):
                    page = pdf_reader.getPage(page_num)
                    content_size += sys.getsizeof(page.extractText())

            elif file_name.endswith(('xlsx', 'xls')):
                if file_name.endswith('xlsx'):
                    sheets = pd.read_excel(BytesIO(object_data), engine='openpyxl', header=None, sheet_name=None)
                else:
                    sheets = pd.read_excel(BytesIO(object_data), engine='xlrd', header=None, sheet_name=None)
                content_size = sys.getsizeof(pd.concat(sheets.values(), ignore_index=True))

            elif file_name.endswith('csv'):
                for delimiter in CSV_DELIMITERS:
                    for encoding in CSV_ENCODINGS:
                        try:
                            content = pd.read_csv(
                                BytesIO(object_data),
                                encoding=encoding,
                                delimiter=delimiter,
                                header=None,
                                engine="python",
                            )
                            return sys.getsizeof(content)
                        except (UnicodeError, pd.errors.ParserError):
                            continue

            elif file_name.endswith('docx'):
                doc = docx.Document(BytesIO(object_data))
                full_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
                content_size = sys.getsizeof('\n'.join(full_text))

            elif file_name.endswith('doc'):
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=True) as doc_temp_file:
                    doc_temp_file.write(object_data)
                    doc_temp_file.flush()  # Ensure data is written to file

                    # Run Antiword to convert the .doc file to text
                    # The output is captured in the subprocess output

                    result = subprocess.run(
                        ['antiword', doc_temp_file.name], check=True, text=True, capture_output=True
                    )
                    content_size = sys.getsizeof(result.stdout)

            elif file_name.endswith(settings.UNSUPPORTED_EXTENSIONS):
                return content_size

            else:
                content_size = sys.getsizeof(object_data.decode('utf-8', errors='ignore'))
        # if raised EmptyDataError it return zero size else we capture error above
        except EmptyDataError:
            logger.warning(f"No columns to parse from file {file_name}")
        return content_size

    def read_archive_object_chunk(
        self,
        chunk_path: str,
        limit: int,
        offset: int,
    ) -> Optional[str]:
        """
        Reads a specific chunk of a file from an archive.

        Args:
            chunk_path: The path of the archive file.
            limit: The size of the chunk to read.
            offset: The offset from where to start reading the chunk.

        Returns:
            The content of the specified chunk as a string, or None if an error occurs.
        """
        file_content = None
        try:
            with open(chunk_path, 'rb') as f:
                file_content = self.prepare_file(f.read(), chunk_path.split('/')[-1], limit, offset)
        except Exception as e:
            logger.error(e)
        return file_content
